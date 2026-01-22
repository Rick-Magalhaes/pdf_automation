from docx import Document

class DocxReader:
    def __init__(self, path: str):
        self.path = path

    def extract_text(self) -> str:
        doc = Document(self.path)
        texts = []

        # 1. Parágrafos normais
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                texts.append(text)

        # 2. Tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        texts.append(cell_text)

        # 3. Cabeçalhos e rodapés
        for section in doc.sections:
            header = section.header
            footer = section.footer

            for p in header.paragraphs:
                text = p.text.strip()
                if text:
                    texts.append(text)

            for p in footer.paragraphs:
                text = p.text.strip()
                if text:
                    texts.append(text)

        return "\n".join(texts)
